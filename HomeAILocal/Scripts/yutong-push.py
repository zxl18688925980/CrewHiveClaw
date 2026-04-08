#!/usr/bin/python3
"""
yutong-push.py
语桐学习数据定时推送到企业微信

用法:
  --mode wrong-questions  推送错题本（默认）
  --mode practice-papers  推送本周练习卷
  --mode all              推送全部
"""
import os
import re
import sys
import argparse
import datetime
from pathlib import Path
from dotenv import load_dotenv
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv('/Users/xinbinanshan/HomeAI/.env')

HOMEAI_ROOT = Path('/Users/xinbinanshan/HomeAI')
FAMILY_DIR = HOMEAI_ROOT / 'Family'
HOMEWORK_DIR = FAMILY_DIR / 'Homework'

CORP_ID = os.environ.get('WECOM_CORP_ID', '')
SECRET = os.environ.get('WECOM_SECRET', '')
AGENT_ID = int(os.environ.get('WECOM_AGENT_ID', '0'))
OWNER_ID = os.environ.get('WECOM_OWNER_ID', '')
MOM_ID = os.environ.get('WECOM_MOM_ID', '')
AUNT_ID = os.environ.get('WECOM_AUNT_ID', '')
FAMILY_CHAT = os.environ.get('WECOM_FAMILY_CHAT_ID', '')


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    lines = md_path.read_text(encoding='utf-8').splitlines()

    def add_run_with_inline(para, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                run = para.add_run(part)

    for line in lines:
        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.runs[0] if h.runs else h.add_run(line[2:])
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            run = h.runs[0] if h.runs else h.add_run(line[3:])
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            add_run_with_inline(p, line[2:])
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x27, 0x6C, 0x2D)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            p.paragraph_format.space_before = Pt(4)
        elif line.strip():
            p = doc.add_paragraph()
            add_run_with_inline(p, line)
            p.paragraph_format.space_after = Pt(2)

    doc.save(docx_path)
    print(f'  Word 文档已生成：{docx_path.name}（{docx_path.stat().st_size//1024}KB）')


def get_token():
    r = requests.get(
        'https://qyapi.weixin.qq.com/cgi-bin/gettoken',
        params={'corpid': CORP_ID, 'corpsecret': SECRET},
        timeout=10
    )
    data = r.json()
    if data.get('errcode', 0) != 0:
        raise RuntimeError(f'获取 token 失败: {data.get("errmsg")}')
    return data['access_token']


def upload_media(token: str, file_path: Path, media_type: str = 'file') -> str:
    mime_types = {
        'file': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf'
    }
    with open(file_path, 'rb') as f:
        r = requests.post(
            f'https://qyapi.weixin.qq.com/cgi-bin/media/upload?access_token={token}&type={media_type}',
            files={'media': (file_path.name, f, mime_types.get(media_type, 'application/octet-stream'))},
            timeout=30
        )
    data = r.json()
    if data.get('errcode', 0) != 0:
        raise RuntimeError(f'上传文件失败: {data.get("errmsg")}')
    return data['media_id']


def send_file_to_user(token: str, media_id: str, to_user: str):
    payload = {
        'touser': to_user,
        'msgtype': 'file',
        'agentid': AGENT_ID,
        'file': {'media_id': media_id}
    }
    r = requests.post(
        f'https://qyapi.weixin.qq.com/cgi-bin/message/send?access_token={token}',
        json=payload, timeout=10
    )
    data = r.json()
    if data.get('errcode', 0) != 0:
        raise RuntimeError(f'发送文件失败: {data.get("errmsg")}')


def send_text_to_user(token: str, text: str, to_user: str):
    payload = {
        'touser': to_user,
        'msgtype': 'text',
        'agentid': AGENT_ID,
        'text': {'content': text}
    }
    r = requests.post(
        f'https://qyapi.weixin.qq.com/cgi-bin/message/send?access_token={token}',
        json=payload, timeout=10
    )
    data = r.json()
    if data.get('errcode', 0) != 0:
        raise RuntimeError(f'发送文字失败: {data.get("errmsg")}')


def send_file_to_group(token: str, media_id: str, chat_id: str):
    payload = {
        'chatid': chat_id,
        'msgtype': 'file',
        'file': {'media_id': media_id}
    }
    r = requests.post(
        f'https://qyapi.weixin.qq.com/cgi-bin/appchat/send?access_token={token}',
        json=payload, timeout=10
    )
    data = r.json()
    if data.get('errcode', 0) != 0:
        raise RuntimeError(f'发送群文件失败: {data.get("errmsg")}')


def push_wrong_questions():
    print('\n[模式 A] 推送英语错题本')
    print('=' * 50)

    md_file = FAMILY_DIR / '英语错题本.md'
    if not md_file.exists():
        print(f'  错题本不存在: {md_file}')
        return False

    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    docx_file = FAMILY_DIR / f'英语错题本_{timestamp}.docx'

    print('\n[1/4] 生成 Word 文档...')
    md_to_docx(md_file, docx_file)

    print('\n[2/4] 获取企业微信 access_token...')
    token = get_token()
    print('  access_token 获取成功')

    print('\n[3/4] 上传文件...')
    media_id = upload_media(token, docx_file)
    print(f'  media_id: {media_id[:20]}...')

    print('\n[4/4] 推送到企业微信...')
    notice = (
        '【Lucas】语桐本周英语错题本已整理完成\n\n'
        'Word 文档见附件，可直接打印使用。'
    )

    for uid in filter(None, [OWNER_ID, MOM_ID, AUNT_ID]):
        send_text_to_user(token, notice, uid)
        send_file_to_user(token, media_id, uid)
        print(f'  已推送给用户: {uid}')

    if FAMILY_CHAT:
        send_file_to_group(token, media_id, FAMILY_CHAT)
        print(f'  已推送到家庭群')

    docx_file.unlink()
    print('\n  错题本推送完成')
    return True


def push_practice_papers():
    print('\n[模式 B] 推送本周练习卷')
    print('=' * 50)

    if not HOMEWORK_DIR.exists():
        print(f'  练习卷目录不存在: {HOMEWORK_DIR}')
        return False

    now = datetime.datetime.now()
    week_ago = now - datetime.timedelta(days=7)

    pdf_files = []
    for f in HOMEWORK_DIR.glob('*.pdf'):
        mtime = datetime.datetime.fromtimestamp(f.stat().st_mtime)
        if mtime >= week_ago:
            pdf_files.append((f, mtime))

    if not pdf_files:
        print('  本周无新增练习卷')
        return True

    pdf_files.sort(key=lambda x: x[1], reverse=True)
    print(f'  发现 {len(pdf_files)} 份本周新增练习卷:')
    for f, mtime in pdf_files:
        print(f'    - {f.name} ({mtime.strftime("%m-%d %H:%M")})')

    print('\n[1/3] 获取企业微信 access_token...')
    token = get_token()
    print('  access_token 获取成功')

    print('\n[2/3] 上传文件...')
    media_ids = []
    for f, _ in pdf_files:
        media_id = upload_media(token, f, 'pdf')
        media_ids.append((f.name, media_id))
        print(f'  上传成功: {f.name}')

    print('\n[3/3] 推送到企业微信...')
    notice = (
        f'【Lucas】语桐本周新增练习卷 {len(pdf_files)} 份\n\n'
        + '\n'.join(f'  - {name}' for name, _ in media_ids)
    )

    for uid in filter(None, [OWNER_ID, MOM_ID, AUNT_ID]):
        send_text_to_user(token, notice, uid)
        for name, media_id in media_ids:
            send_file_to_user(token, media_id, uid)
        print(f'  已推送给用户: {uid}')

    if FAMILY_CHAT:
        for name, media_id in media_ids:
            send_file_to_group(token, media_id, FAMILY_CHAT)
        print(f'  已推送到家庭群')

    print('\n  练习卷推送完成')
    return True


def main():
    parser = argparse.ArgumentParser(description='语桐学习数据定时推送')
    parser.add_argument('--mode', choices=['wrong-questions', 'practice-papers', 'all'],
                        default='wrong-questions', help='推送模式')
    args = parser.parse_args()

    if not CORP_ID or not SECRET:
        print('错误: 缺少 WECOM_CORP_ID 或 WECOM_SECRET 环境变量')
        sys.exit(1)

    print('\n语桐学习数据推送')
    print(f'时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    print(f'模式: {args.mode}')

    if args.mode == 'wrong-questions':
        push_wrong_questions()
    elif args.mode == 'practice-papers':
        push_practice_papers()
    elif args.mode == 'all':
        push_wrong_questions()
        push_practice_papers()

    print('\n全部完成')


if __name__ == '__main__':
    main()
